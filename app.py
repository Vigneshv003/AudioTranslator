import os
import tempfile
import streamlit as st
from faster_whisper import WhisperModel
from docx import Document

# ==========================================================
# PAGE CONFIG
# ==========================================================

st.set_page_config(
    page_title="Audio Translator",
    page_icon="🎙️",
    layout="centered"
)

st.title("🎙️ Spanish Audio → English Word")
st.write("Upload a WAV, MP3, or M4A file and download the English transcript.")

# ==========================================================
# SETTINGS
# ==========================================================

MAX_FILE_SIZE = 25 * 1024 * 1024  # 25 MB

# ==========================================================
# LOAD MODEL (Cached)
# ==========================================================

@st.cache_resource(show_spinner=False)
def load_model(model_name):
    return WhisperModel(
        model_name,
        device="cpu",
        compute_type="int8"
    )

# ==========================================================
# MODEL SELECTION
# ==========================================================

MODEL_NAME = st.selectbox(
    "Select Whisper Model",
    ["tiny", "base", "small"],
    index=1  # base recommended
)

st.caption("Recommended: **base** for best speed and accuracy on Streamlit Cloud.")

# ==========================================================
# FILE UPLOAD
# ==========================================================

uploaded_file = st.file_uploader(
    "Upload Audio",
    type=["wav", "mp3", "m4a"]
)

# ==========================================================
# MAIN
# ==========================================================

if uploaded_file:

    # Check file size
    if uploaded_file.size > MAX_FILE_SIZE:
        st.error("❌ File size exceeds 25 MB.")
        st.stop()

    st.write(f"**File:** {uploaded_file.name}")
    st.write(f"**Size:** {uploaded_file.size / 1024 / 1024:.2f} MB")

    if st.button("Convert to Word", type="primary"):

        temp_audio_path = None
        output_doc_path = None

        try:
            progress = st.progress(0)

            # Load model
            progress.progress(10)
            with st.spinner("Loading Whisper model..."):
                model = load_model(MODEL_NAME)

            # Save uploaded audio
            progress.progress(30)

            suffix = os.path.splitext(uploaded_file.name)[1]

            with tempfile.NamedTemporaryFile(delete=False, suffix=suffix) as temp_audio:
                temp_audio.write(uploaded_file.getbuffer())
                temp_audio_path = temp_audio.name

            # Transcribe & translate
            progress.progress(50)

            with st.spinner("Transcribing and translating audio..."):

                segments, info = model.transcribe(
                    temp_audio_path,
                    task="translate",
                    language="es"
                )

                transcript = " ".join(segment.text for segment in segments)

            progress.progress(80)

            # Create Word document
            doc = Document()
            doc.add_heading("Meeting Transcript", level=1)
            doc.add_paragraph(transcript)

            output_doc_path = tempfile.NamedTemporaryFile(
                delete=False,
                suffix=".docx"
            ).name

            doc.save(output_doc_path)

            progress.progress(100)

            st.success("✅ Translation completed successfully!")

            # Preview
            with st.expander("📄 Preview Transcript"):
                st.write(transcript[:3000] + ("..." if len(transcript) > 3000 else ""))

            # Download
            with open(output_doc_path, "rb") as f:
                st.download_button(
                    label="📥 Download Word File",
                    data=f.read(),
                    file_name="Meeting_Transcript.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                )

        except Exception as e:
            st.error(f"❌ Error: {str(e)}")

        finally:
            # Cleanup temp files
            if temp_audio_path and os.path.exists(temp_audio_path):
                os.remove(temp_audio_path)

            if output_doc_path and os.path.exists(output_doc_path):
                os.remove(output_doc_path)
